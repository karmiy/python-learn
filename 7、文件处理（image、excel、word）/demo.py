"""image"""

# pip install pillow

# from PIL import Image
# import os

# image = Image.open(os.path.join(os.path.dirname(__file__), "suzy.png"))
# image.format, image.size, image.mode
# ("JPEG", (500, 750), "RGB")
# image.rotate(180).show()  # 旋转

"""Excel"""
# pip install openpyxl

# import datetime
# import os

# from openpyxl import Workbook

# wb = Workbook()
# ws = wb.active

# ws["A1"] = 42
# ws.append([1, 2, 3])
# ws["A2"] = datetime.datetime.now()

# wb.save(os.path.join(os.path.dirname(__file__), "sample.xlsx"))

"""Word"""
# pip install python-docx

from docx import Document
from docx.shared import Inches
import os

document = Document()

document.add_heading("Document Title", 0)

p = document.add_paragraph("A plain paragraph having some ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic.").italic = True

document.add_heading("Heading, level 1", level=1)
document.add_paragraph("Intense quote", style="Intense Quote")

document.add_paragraph("first item in unordered list", style="List Bullet")
document.add_paragraph("first item in ordered list", style="List Number")

document.add_picture(
    os.path.join(os.path.dirname(__file__), "suzy.png"), width=Inches(1.25)
)

records = (
    (3, "101", "Spam"),
    (7, "422", "Eggs"),
    (4, "631", "Spam, spam, eggs, and spam"),
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "Id"
hdr_cells[2].text = "Desc"
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save(os.path.join(os.path.dirname(__file__), "demo.docx"))
